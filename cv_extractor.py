#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
CV Information Extractor (Gemini 2.5 Flash)
-------------------------------------------
Reads a CV (PDF or DOCX), extracts text, sends it to Gemini API,
and returns a structured JSON + ATS Word file.

Environment variable required:
    GEMINI_API_KEY=...
"""

import os
import re
import json
import sys
from datetime import datetime
from typing import Dict, Any, List

import fitz  # PyMuPDF for PDF text
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()

# ---------------- CONFIG ----------------
MODEL_ID = "gemini-2.5-flash"
TEMPERATURE = 0.0
TIMEOUT_S = 120

JSON_SCHEMA = {
    "full_name": None,
    "current_title": None,
    "total_experience_years": None,
    "areas_of_expertise": [],
    "industry_expertise": [],
    "education": [],
    "professional_memberships": [],
    "trainings_and_certifications": [],
    "professional_experience": [],
    "projects": [],
}


# ----------- NORMALIZATION HELPERS -----------
def _as_list(x):
    if x is None:
        return []
    return x if isinstance(x, list) else [x]


def _coerce_str_list(items) -> List[str]:
    return [str(i).strip() for i in _as_list(items) if str(i).strip()]


def _dedupe_preserve(seq: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in seq:
        if x not in seen and str(x).strip():
            seen.add(x)
            out.append(x)
    return out


def _smart_title(s: str) -> str:
    parts = []
    for w in s.split():
        if len(w) <= 3 or w.isupper():
            parts.append(w.upper())
        else:
            parts.append(w.capitalize())
    return " ".join(parts)


def _normalize_skills(items) -> List[str]:
    vals = _coerce_str_list(items)
    vals = _dedupe_preserve(vals)
    return [_smart_title(v) for v in vals]


def _coerce_education(items) -> List[Dict[str, Any]]:
    out = []
    for ed in _as_list(items):
        if isinstance(ed, dict):
            out.append(
                {
                    "degree": ed.get("degree"),
                    "institution": ed.get("institution"),
                    "year": ed.get("year"),
                }
            )
        else:
            s = str(ed).strip()
            if not s:
                continue
            parts = [p.strip() for p in s.split("|")]
            degree = parts[0] if len(parts) > 0 else None
            institution = parts[1] if len(parts) > 1 else None
            year = parts[2] if len(parts) > 2 else None
            out.append({"degree": degree, "institution": institution, "year": year})
    return out


def _coerce_experience(items) -> List[Dict[str, Any]]:
    out = []
    for exp in _as_list(items):
        if isinstance(exp, dict):
            out.append(
                {
                    "company": exp.get("company"),
                    "title": exp.get("title"),
                    "dates": exp.get("dates"),
                    "highlights": _coerce_str_list(exp.get("highlights")),
                }
            )
        else:
            s = str(exp).strip()
            if not s:
                continue
            out.append(
                {
                    "company": None,
                    "title": None,
                    "dates": s,
                    "highlights": [],
                }
            )
    return out


def _coerce_projects(items) -> List[Dict[str, Any]]:
    out = []
    for pr in _as_list(items):
        if isinstance(pr, dict):
            out.append(
                {
                    "name": pr.get("name"),
                    "role": pr.get("role"),
                    "summary": pr.get("summary"),
                    "tech": _coerce_str_list(pr.get("tech")),
                }
            )
        else:
            s = str(pr).strip()
            if not s:
                continue
            out.append(
                {
                    "name": s,
                    "role": None,
                    "summary": None,
                    "tech": [],
                }
            )
    return out


def _coerce_certifications(items):
    out = []
    for c in _as_list(items):
        if isinstance(c, dict):
            name = str(c.get("name") or "").strip()
            issuer = str(c.get("issuer") or "").strip()
            duration = str(c.get("duration") or "").strip()
            out.append(
                {
                    "name": name or None,
                    "issuer": issuer or None,
                    "duration": duration or None,
                }
            )
        else:
            s = str(c).strip()
            if s:
                out.append(s)
    return out


def _coerce_years(x):
    if x is None:
        return None
    try:
        return float(str(x).split()[0])
    except Exception:
        return None


# ---- Experience years inference helpers ----
MONTHS = {
    m.lower(): i
    for i, m in enumerate(
        [
            "Jan",
            "Feb",
            "Mar",
            "Apr",
            "May",
            "Jun",
            "Jul",
            "Aug",
            "Sep",
            "Oct",
            "Nov",
            "Dec",
        ],
        start=1,
    )
}
FULL_MONTHS = {
    m.lower(): i
    for i, m in enumerate(
        [
            "January",
            "February",
            "March",
            "April",
            "May",
            "June",
            "July",
            "August",
            "September",
            "October",
            "November",
            "December",
        ],
        start=1,
    )
}


def _parse_date_token(tok: str):
    tok = tok.strip().strip(".").replace(",", "")
    if tok.lower() in {"present", "current", "now"}:
        today = datetime.today()
        return datetime(today.year, today.month, 1)
    parts = tok.split()
    if len(parts) == 2:
        m, y = parts[0].lower(), parts[1]
        if y.isdigit():
            if m in MONTHS:
                return datetime(int(y), MONTHS[m], 1)
            if m in FULL_MONTHS:
                return datetime(int(y), FULL_MONTHS[m], 1)
    m = re.match(r"(\d{4})[-/](\d{1,2})", tok)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), 1)
    m = re.match(r"(\d{4})$", tok)
    if m:
        return datetime(int(m.group(1)), 1, 1)
    return None


def _infer_years_from_experience(exps) -> float | None:
    if not exps:
        return None
    months_total = 0
    any_found = False
    for exp in exps:
        dates = (
            (exp.get("dates") or "").strip()
            if isinstance(exp, dict)
            else str(exp).strip()
        )
        if not dates:
            continue
        norm = dates.replace("—", "-").replace("–", "-").replace("to", "-")
        if "-" in norm:
            parts = [p.strip() for p in norm.split("-", 1)]
            if len(parts) == 2:
                start = _parse_date_token(parts[0])
                end = _parse_date_token(parts[1])
                if start:
                    if not end:
                        end = datetime.today()
                    if end >= start:
                        any_found = True
                        months_total += (end.year - start.year) * 12 + (
                            end.month - start.month
                        )
        else:
            dt = _parse_date_token(norm)
            if dt:
                any_found = True
                months_total += 12
    if any_found and months_total > 0:
        return round(months_total / 12.0, 1)
    return None


# ---- Projects heuristic (fallback) ----
PROJECT_HEADINGS = re.compile(
    r"(?im)^\s*(projects?|selected projects?|academic projects?)\s*$"
)


def _find_projects_block(cv_text: str) -> str | None:
    m = PROJECT_HEADINGS.search(cv_text)
    if not m:
        return None
    start = m.end()
    NEXT_HEADING = re.compile(
        r"(?im)^\s*([A-Z][A-Z \-/&]{3,}|[A-Z][a-z]+(?: [A-Z][a-z]+){0,3}\s*:|[0-9]+\.\s+[A-Z].+)\s*$"
    )
    tail = cv_text[start:]
    n = NEXT_HEADING.search(tail)
    end = start + n.start() if n else len(cv_text)
    block = cv_text[start:end].strip()
    return block if block else None


def heuristic_projects_from_text(cv_text: str, max_items: int = 5):
    block = _find_projects_block(cv_text)
    if not block:
        return []
    lines = [ln.strip(" \t•*-") for ln in block.splitlines()]
    lines = [
        ln
        for ln in lines
        if ln and not ln.lower().startswith(("references", "referees"))
    ]

    projects = []
    current = {"name": None, "role": None, "summary": None, "tech": []}

    def flush():
        if current["name"] or current["summary"]:
            current["name"] = (current["name"] or "").strip() or None
            current["summary"] = (current["summary"] or "").strip() or None
            current["tech"] = [t.strip() for t in current["tech"] if t.strip()]
            projects.append(current.copy())

    TECH_PREFIXES = ("tech:", "tools:", "stack:", "technologies:", "technology:")

    for ln in lines:
        lower = ln.lower()
        if any(lower.startswith(p) for p in TECH_PREFIXES):
            tech_str = ln.split(":", 1)[1] if ":" in ln else ""
            techs = [t.strip() for t in tech_str.split(",") if t.strip()]
            current["tech"].extend(techs)
            continue

        if re.search(r"[:—–-]\s*$", ln) or len(ln.split()) <= 8:
            if current["name"] or current["summary"] or current["tech"]:
                flush()
                current = {"name": None, "role": None, "summary": None, "tech": []}
            if ln.endswith(":"):
                current["name"] = ln[:-1].strip()
            else:
                if len(ln.split()) <= 8:
                    current["name"] = ln
                else:
                    current["summary"] = ln
            continue

        if current["summary"]:
            current["summary"] += " " + ln
        else:
            current["summary"] = ln

    flush()
    cleaned = []
    for pr in projects:
        if not pr.get("name") and pr.get("summary"):
            words = pr["summary"].split()
            pr["name"] = " ".join(words[:6]) + ("..." if len(words) > 6 else "")
        cleaned.append(
            {
                "name": pr.get("name"),
                "role": pr.get("role"),
                "summary": pr.get("summary"),
                "tech": pr.get("tech") or [],
            }
        )
    return cleaned[:max_items]


def normalize_parsed(data: dict) -> dict:
    areas = _normalize_skills(data.get("areas_of_expertise"))
    industries = _normalize_skills(data.get("industry_expertise"))
    education = _coerce_education(data.get("education"))
    experience = _coerce_experience(data.get("professional_experience"))
    projects = _coerce_projects(data.get("projects"))
    certs = _coerce_certifications(data.get("trainings_and_certifications"))

    years = _coerce_years(data.get("total_experience_years"))
    if years is None:
        inferred = _infer_years_from_experience(experience)
        if inferred is not None:
            years = inferred

    if years is not None:
        years = int(round(years))

    return {
        "full_name": data.get("full_name"),
        "current_title": data.get("current_title"),
        "total_experience_years": years,
        "areas_of_expertise": areas,
        "industry_expertise": industries,
        "education": education,
        "professional_memberships": _coerce_str_list(
            data.get("professional_memberships")
        ),
        "trainings_and_certifications": certs,
        "professional_experience": experience,
        "projects": projects,
    }


# --------------- UTILITIES ---------------
def read_cv_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_text_from_pdf(path)
    elif ext == ".docx":
        return _extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_text_from_pdf(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    texts = [page.get_text("text") for page in doc]
    doc.close()
    return "\n".join(texts).strip()


def _extract_text_from_docx(docx_path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


# --------------- PROMPT ------------------
def build_prompt(cv_text: str) -> str:
    schema_str = json.dumps(JSON_SCHEMA, indent=2, ensure_ascii=False)
    return f"""
You are an expert ATS parser. Extract ONLY the following fields from the CV text
and return a SINGLE JSON object with EXACTLY these keys:

{schema_str}

Rules:
- Use null for missing scalar fields and [] for empty lists.
- "total_experience_years" must be numeric or null.
- Keep original wording/spelling; do NOT invent data.
- If a *Projects* section (or similar: "Selected Projects", "Academic Projects") exists, extract up to 5 entries.
- For each project, try to fill:
  - name: short title (line before colon or bold phrase)
  - role: if explicitly stated; otherwise null
  - summary: 1–2 short phrases from bullet(s)
  - tech: list parsed from "Tech/Tools/Stack" lines (split by comma)
- If projects are only bullets without names, use the first 5–7 words as a short name and the full bullet as summary.
- Output MUST be pure JSON (no commentary), exactly the schema keys.

CV TEXT:
{cv_text}
"""


# --------------- GEMINI CALL --------------
def extract_with_gemini(cv_text: str) -> Dict[str, Any]:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("❌ Environment variable GEMINI_API_KEY not set!")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name=MODEL_ID,
        generation_config={
            "temperature": TEMPERATURE,
            "response_mime_type": "application/json",
        },
    )

    response = model.generate_content(
        build_prompt(cv_text), request_options={"timeout": TIMEOUT_S}
    )

    text = (getattr(response, "text", None) or "").strip()
    if not text and getattr(response, "candidates", None):
        try:
            text = response.candidates[0].content.parts[0].text.strip()
        except Exception:
            pass

    if not text:
        raise RuntimeError("Gemini returned an empty response.")

    try:
        data = json.loads(text)
    except Exception:
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end != -1:
            data = json.loads(text[start : end + 1])
        else:
            raise RuntimeError("Gemini output not valid JSON:\n" + text)

    data = {k: data.get(k, JSON_SCHEMA[k]) for k in JSON_SCHEMA}
    return data


# --------------- WORD OUTPUT --------------
def _remove_table_borders(table):
    """
    Robustly remove all borders from a python-docx table by editing the XML.
    Avoids direct attribute access that may not exist across versions.
    """
    tbl = table._tbl  # CT_Tbl (lxml element)
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # find or create <w:tblBorders>
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # for each border edge ensure w:val="nil"
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag_qn = qn(f'w:{edge}')
        el = tblBorders.find(tag_qn)
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tblBorders.append(el)
        el.set(qn('w:val'), 'nil')


def write_ats_docx(parsed: Dict[str, Any], out_path: str):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def space(lines=1):
        for _ in range(lines):
            doc.add_paragraph()

    def heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)

    # ===== Header =====
    name = parsed.get("full_name") or ""
    title = parsed.get("current_title") or ""
    head = doc.add_paragraph()
    r1 = head.add_run(name)
    r1.bold = True
    r1.font.size = Pt(16)
    if title:
        head.add_run(" — " + title)
    space(1)

    # ===== Total years =====
    if parsed.get("total_experience_years") is not None:
        doc.add_paragraph(f"Total Experience: {parsed['total_experience_years']} years")

    # ===== Professional Experience (moved to the top) =====
    if parsed.get("professional_experience"):
        space(1)
        heading("Professional Experience")
        for exp in parsed["professional_experience"]:
            if isinstance(exp, dict):
                head_parts = [exp.get("title"), exp.get("company"), exp.get("dates")]
                head_line = " | ".join([p for p in head_parts if p])
                if head_line:
                    doc.add_paragraph(head_line)
                for h in exp.get("highlights", []):
                    if h and str(h).strip():
                        doc.add_paragraph(str(h).strip(), style="List Bullet")
            else:
                line = str(exp).strip()
                if line:
                    doc.add_paragraph(line)

    # ===== Areas of Expertise (borderless 2-column table) =====
    areas = parsed.get("areas_of_expertise") or []
    if areas:
        space(1)
        heading("Areas of Expertise")

        mid = (len(areas) + 1) // 2
        left_items, right_items = areas[:mid], areas[mid:]

        table = doc.add_table(rows=1, cols=2)
        _remove_table_borders(table)  # remove borders reliably

        # Left column
        left_cell = table.rows[0].cells[0]
        for item in left_items:
            p = left_cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(item))

        # Right column
        right_cell = table.rows[0].cells[1]
        for item in right_items:
            p = right_cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(item))

    # ===== Industry Expertise =====
    industries = parsed.get("industry_expertise") or []
    if industries:
        space(1)
        heading("Industry Expertise")
        for item in industries:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(item))

    # ===== Education =====
    education = parsed.get("education") or []
    if education:
        space(1)
        heading("Education")
        for ed in education:
            if isinstance(ed, dict):
                parts = [ed.get("degree"), ed.get("institution"), ed.get("year")]
                parts = [p for p in parts if p]
                line = " | ".join(parts)
            else:
                line = str(ed).strip()
            if line:
                doc.add_paragraph(line)

    # ===== Professional Memberships =====
    memberships = parsed.get("professional_memberships") or []
    if memberships:
        space(1)
        heading("Professional Memberships")
        for m in memberships:
            doc.add_paragraph(str(m), style="List Bullet")

    # ===== Trainings & Certifications =====
    certs = parsed.get("trainings_and_certifications") or []
    if certs:
        space(1)
        heading("Trainings & Certifications")
        for c in certs:
            if isinstance(c, dict):
                parts = []
                if c.get("name"):
                    parts.append(c["name"])
                if c.get("issuer"):
                    parts.append(c["issuer"])
                line = " — ".join(parts) if parts else ""
                if c.get("duration"):
                    line = (line + f" ({c['duration']})").strip()
                if line:
                    doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(str(c), style="List Bullet")

    # ===== Projects =====
    projects = parsed.get("projects") or []
    if projects:
        space(1)
        heading("Projects")
        for pr in projects:
            if isinstance(pr, dict):
                header = " | ".join([p for p in [pr.get("name"), pr.get("role")] if p])
                if header:
                    doc.add_paragraph(header)
                if pr.get("summary"):
                    doc.add_paragraph(pr["summary"])
                if pr.get("tech"):
                    doc.add_paragraph("Technologies: " + ", ".join(pr["tech"]))
            else:
                line = str(pr).strip()
                if line:
                    doc.add_paragraph(line)

    doc.save(out_path)


# --------------- MAIN ---------------------
def main():
    if len(sys.argv) < 2:
        print(
            "Usage: python cv_extractor.py <path_to_cv.(pdf|docx)> [--out out.docx] [--json out.json]"
        )
        sys.exit(1)

    cv_path = sys.argv[1]
    out_docx = "cv_parsed_ats.docx"
    out_json = "cv_parsed.json"

    if "--out" in sys.argv:
        out_docx = sys.argv[sys.argv.index("--out") + 1]
    if "--json" in sys.argv:
        out_json = sys.argv[sys.argv.index("--json") + 1]

    print(f"[1/3] Reading CV: {cv_path}")
    text = read_cv_text(cv_path)

    print(f"[2/3] Extracting with Gemini ({MODEL_ID}) ...")
    data = extract_with_gemini(text)
    data = normalize_parsed(data)

    # 🔁 Fallback: mine Projects from text if empty
    if not data.get("projects"):
        mined = heuristic_projects_from_text(text, max_items=5)
        if mined:
            data["projects"] = mined

    with open(out_json, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"   → JSON saved to {out_json}")

    print("[3/3] Writing ATS Word file...")
    write_ats_docx(data, out_docx)
    print(f"   → DOCX saved to {out_docx}")
    print("\n✅ Done!")


if __name__ == "__main__":
    main()
